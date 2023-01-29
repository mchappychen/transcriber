import time, os, re, subprocess, datetime
import numpy as np
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from send2trash import send2trash #I use this to send .wma to trash
from mutagen.mp3 import MP3
from mutagen.mp4 import MP4
from mutagen.asf import ASF
from mutagen.wave import WAVE
import nltk

model_name = 'base.en'
#tiny.en is fast but bad, use only for testing
#base.en is 1.5x slower but better, so I use this
#small.en is super slow, maybe even slower while using otranscribe at the same time
#medium.en cannot be handled by my 4GB RAM laptop

#pink,red,yellow,green,blue,violet,end
class c:
   violet = '\033[95m'
   pink = '\033[94m'
   blue = '\033[96m'
   green = '\033[92m'
   yellow = '\033[93m'
   red = '\033[91m'
   end = '\033[0m'
   bold = '\033[1m'

#importing whisper takes 20s
print(f'{c.blue}Importing whisper ... (~20s){c.end}')
import whisper

print(f'{c.blue}loading model {model_name} ... (~2s){c.end}')
model = whisper.load_model(model_name)

data = open('data.csv','a')

#Prediction models for how long a model will take based on previous data
def predict(length):
	return round(0.288*length + 2.225)

#returns the length of the file
def get_length(filename):
	try:
		if filename[-3:].lower() == 'mp3':
			return round(MP3(filename).info.length)
		elif filename[-3:].lower() == 'm4a':
			return round(MP4(filename).info.length)
		elif filename[-3:].lower() == 'wma':
			return round(ASF(filename).info.length)
		elif filename[-3:].lower() == 'wav':
			return round(WAVE(filename).info.length)
	except Exception as e:
		print(f'{c.red}{filename} Error:{e}{c.blue}')
	return None

#Extracts names from text
def get_human_names(text):
	try:
		person_list = []
		tokens = nltk.tokenize.word_tokenize(text)
		pos = nltk.pos_tag(tokens)
		sentt = nltk.ne_chunk(pos, binary = False)
		person = []
		name = ""
		for subtree in sentt.subtrees(filter=lambda t: t.label() == 'PERSON'):
			for leaf in subtree.leaves():
				person.append(leaf[0])
			if len(person) > 1: #avoid grabbing lone surnames
				for part in person:
					name += part + ' '
				if name[:-1] not in person_list:
					person_list.append(name[:-1])
				name = ''
			person = []
	except LookupError:
		print(f'{c.yellow}Downloading nltk stuff (one time only)...{c.blue}')
		nltk.download('punkt')
		nltk.download('averaged_perceptron_tagger')
		nltk.download('maxent_ne_chunker')
		nltk.download('words')
		nltk.download('omw-1.4')
		person_list = get_human_names(text)
	return person_list

while input(f'{c.green}Enter any key to start next batch...{c.blue}'):

	#get all mp3 files
	files = [f for f in os.listdir('.') if os.path.isfile(f) and f[-4:].lower() in ['.mp3','.m4a','.wma','.wav']]

	#transcribe each file
	for file in files:

		length = get_length(file)
		print(f'{c.blue}Transcribing: {c.pink}{file} {c.blue}Length: {c.pink}{length//60}:{length%60}{c.blue}... (~{c.pink}{predict(length)}s{c.blue}){c.end}')

		#Convert .wma to .mp3 (for otranscribe)
		if file[-4:].lower() == '.wma':
			print(f'{c.blue}Converting into .mp3 in background ...')
			os.system(f'ffmpeg -i '+ file.replace(' ','\ ') + ' '+ file[:-4].replace(' ','\ ') +'.mp3 -hide_banner -loglevel panic')

		start = time.time()
		#transcribe
		result = model.transcribe(file,fp16=False,language='English',temperature=0)

		#Gets rid of the weird space in the beginning
		output = result['text'][1:]

		#Add content to word doc
		document = Document()
		document.add_paragraph(f'{file[:-4]}\n\nTranscriber - Michael Chen\n')
		texts = []

		#Format output better
		output = output.split('? ')
		for x in range(len(output)):
			text = output[x]
			if x < len(output)-1:
				text += '?'
				if x > 0:
					text = '.\n\n'.join(text.rsplit('. ',1))
			#Find names
			if 'name' in text:
				names = get_human_names(text)
				if len(names):
					#if there are names, add it to beginning
					for name in names:
						document.add_paragraph(f'(man) {name} - \n')
			texts.append(text)

		for text in texts:
			document.add_paragraph(text)

		#style the document
		for paragraph in document.paragraphs:
			paragraph.style = document.styles['Normal']
			for run in paragraph.runs:
				run.font.size = Pt(16)
				run.font.name = 'Arial'

		document.save(f'{file[:-4]}.docx')

		print(f'{c.blue}Done. Time took: {c.green}{round(time.time()-start)}s{c.end}')
		data.write(f'\n{get_length(file)},{round(time.time()-start)},{file[-3:]}')
		data.flush() #flush the buffer from data.write into the file
		# subprocess.call(['open',file[:-4]+'.docx']) #this auto-opens a file

	#trash all the .wma files
	for file in files:
		if file[-4:].lower() == '.wma':
			print(f'{c.blue}Moving {c.pink}{file}{c.blue} to trash{c.end}')
			send2trash(file)


